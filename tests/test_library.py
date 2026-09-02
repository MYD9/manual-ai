import hashlib
import io
import json
import os
import sqlite3
import tempfile
import zipfile
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path

os.environ["MANUAL_AI_HOME"] = tempfile.mkdtemp(prefix="manual-ai-test-bootstrap-")
os.environ["MANUAL_AI_NO_WORKER"] = "1"
import pytest
import numpy as np
from fastapi.testclient import TestClient
from sqlalchemy import select
from backend import main
from backend.models import Chunk, Entry, Job, Vector
from backend.processing import Processor, parse_source, safe_url
from backend.storage import Library, protect
from backend.providers import space_for


@pytest.fixture
def client(tmp_path, monkeypatch):
    lib = Library(tmp_path / "data")
    monkeypatch.setattr(main, "library", lib)
    monkeypatch.setattr(main, "processor", Processor(lib))
    with TestClient(main.app) as client:
        yield client
    lib.engine.dispose()


def create(client, **data):
    r = client.post('/api/v1/entries', json={"title": "ESP32-S3", **data})
    assert r.status_code == 201, r.text
    return r.json()


def search(client, q, **data):
    r = client.post('/api/v1/search', json={"q": q, "mode": "keyword", **data})
    assert r.status_code == 200, r.text
    return r.json()["items"]


def process_upload(client, mid, name, data, mime='application/octet-stream'):
    r = client.post('/api/v1/imports/file', data={"manual_id": mid}, files={"file": (name, data, mime)})
    assert r.status_code == 202, r.text
    value = r.json()
    main.processor.process(value['job_id'], value['source']['id'])
    return client.get('/api/v1/entries/' + value['source']['id']).json()


def test_persistence_search_revision_and_html(client):
    m = create(client)
    n = create(client, kind='note', manual_id=m['id'], title='供电与串口', content='<p>工作电压 3.3V，GPIO 18 可用于示例。中文搜索。</p><script>alert(1)</script><img src="https://tracker.invalid/a">')
    assert '<script' not in n['content'] and 'tracker.invalid' not in n['content']
    assert any(x['entry_id'] == n['id'] for x in search(client, '电压'))
    assert any(x['entry_id'] == n['id'] for x in search(client, 'GPIO18'))
    url='/api/v1/entries/' + n['id']
    assert client.patch(url, json={'revision': 1, 'content': '<p>新的内容</p>'}).status_code == 200
    assert client.patch(url, json={'revision': 1, 'content': '<p>旧草稿</p>'}).status_code == 409
    assert not search(client, 'GPIO18')
    reopened = Library(main.library.root)
    with reopened.Session() as db:
        assert db.get(Entry,n['id']).content == '<p>新的内容</p>'
    reopened.engine.dispose()


def test_two_simultaneous_edits_only_one_wins(client):
    m=create(client)
    def change(title):
        return client.patch('/api/v1/entries/'+m['id'],json={'revision':1,'title':title}).status_code
    with ThreadPoolExecutor(max_workers=2) as pool:
        statuses=list(pool.map(change,['甲','乙']))
    assert sorted(statuses)==[200,409]


def test_nested_trash_restores_only_matching_batch(client):
    m=create(client)
    chapter=create(client,kind='chapter',manual_id=m['id'],title='章节')
    a=create(client,kind='card',manual_id=m['id'],parent_id=chapter['id'],title='电压资料',content='3.3V')
    b=create(client,kind='note',manual_id=m['id'],title='先删除的笔记')
    assert client.post('/api/v1/entries/'+b['id']+'/trash',json={'revision':1}).status_code==200
    assert client.post('/api/v1/entries/'+m['id']+'/trash',json={'revision':1}).status_code==200
    assert not search(client,'电压')
    assert client.post('/api/v1/entries/'+a['id']+'/restore').status_code==409
    assert client.post('/api/v1/entries/'+m['id']+'/restore').status_code==200
    assert client.get('/api/v1/entries/'+a['id']).status_code==200
    assert client.get('/api/v1/entries/'+b['id']).status_code==404
    assert search(client,'电压')


def test_move_and_merge_preserve_cards(client):
    m=create(client); other=create(client,title='另一说明书')
    a=create(client,kind='chapter',manual_id=m['id'],title='甲')
    b=create(client,kind='chapter',manual_id=m['id'],title='乙')
    card=create(client,kind='card',manual_id=m['id'],parent_id=a['id'],title='引脚说明')
    r=client.post('/api/v1/chapters/merge',json={'source_id':a['id'],'target_id':b['id'],'source_revision':1,'target_revision':1})
    assert r.status_code==200,r.text
    assert client.get('/api/v1/entries/'+card['id']).json()['parent_id']==b['id']
    assert client.patch('/api/v1/entries/'+b['id'],json={'revision':2,'manual_id':other['id']}).status_code==200
    moved=client.get('/api/v1/entries/'+card['id']).json()
    assert moved['manual_id']==other['id']
    assert not search(client,'引脚',manual_id=m['id'])
    assert search(client,'引脚',manual_id=other['id'])


def test_duplicate_blob_not_deleted_with_one_source(client):
    a=create(client); b=create(client,title='第二本')
    raw='共享原件：UART 引脚设置。'.encode()
    one=process_upload(client,a['id'],'说明.txt',raw)
    two=process_upload(client,b['id'],'说明.txt',raw)
    assert one['attrs']['hash']==two['attrs']['hash']
    r=client.post('/api/v1/imports/file',data={'manual_id':a['id']},files={'file':('说明.txt',raw)})
    assert r.json()['duplicate']
    client.post('/api/v1/entries/'+one['id']+'/trash',json={'revision':one['revision']})
    assert client.get('/api/v1/blobs/'+two['attrs']['hash']).content==raw
    assert all(h['entry_id']!=one['id'] for h in search(client,'UART'))


def test_docx_tables_images_and_corrupt_pdf(client):
    import docx
    from PIL import Image
    doc=docx.Document();doc.add_heading('测试设备说明',0);doc.add_paragraph('供电电压 5V，USB 连接。');table=doc.add_table(rows=2,cols=2);table.cell(0,0).text='参数';table.cell(1,0).text='电压';table.cell(1,1).text='5V'
    image=io.BytesIO();Image.new('RGB',(20,20),'blue').save(image,format='PNG');image.seek(0);doc.add_picture(image)
    data=io.BytesIO();doc.save(data)
    m=create(client)
    source=process_upload(client,m['id'],'测试.docx',data.getvalue())
    assert len(source['attrs']['images'])==1
    assert any('5V' in b['text'] for b in source['attrs']['blocks'])
    assert search(client,'电压')
    r=client.post('/api/v1/imports/file',data={'manual_id':m['id']},files={'file':('bad.pdf',b'not a pdf')})
    with pytest.raises(ValueError,match='PDF'):
        main.processor.process(r.json()['job_id'],r.json()['source']['id'])


def test_vector_filtering_deleted_disabled_and_changed(client,monkeypatch):
    m=create(client,attrs={'ai_enabled':True})
    c=create(client,kind='card',manual_id=m['id'],title='限制',content='引脚只能输入')
    space=space_for({})
    with main.library.Session() as db:
        chunks=db.scalars(select(Chunk)).all()
        for chunk in chunks:
            value=np.array([1,0,0],dtype=np.float32)
            db.add(Vector(chunk_id=chunk.id,space=space,digest=chunk.digest,dimensions=3,value=value.tobytes()))
        db.commit()
    monkeypatch.setattr(main,'embeddings',lambda *args:[[1,0,0]])
    r=client.post('/api/v1/search',json={'q':'输入限制','mode':'semantic'})
    assert c['id'] in [i['entry_id'] for i in r.json()['items']]
    client.post('/api/v1/entries/'+c['id']+'/trash',json={'revision':1})
    assert c['id'] not in [i['entry_id'] for i in client.post('/api/v1/search',json={'q':'输入限制','mode':'semantic'}).json()['items']]
    client.patch('/api/v1/entries/'+m['id'],json={'revision':1,'attrs':{'ai_enabled':False}})
    assert client.post('/api/v1/search',json={'q':'限制','mode':'semantic'}).json()['items']==[]


def test_sse_citations_and_no_credentials_in_response(client,monkeypatch):
    m=create(client,attrs={'ai_enabled':True})
    create(client,kind='card',manual_id=m['id'],title='测试电压',content='电压是 3.3V')
    monkeypatch.setattr(main,'api_config',lambda *args:('https://example.invalid','test','test'))
    monkeypatch.setattr(main,'embeddings',lambda *args:(_ for _ in ()).throw(ValueError('未配置向量')))
    monkeypatch.setattr(main,'chat_stream',lambda *args:iter(['根据资料，','电压是 3.3V [1]。']))
    r=client.post('/api/v1/chat',json={'question':'电压是多少','manual_id':m['id']})
    assert 'event: sources' in r.text and 'event: delta' in r.text and 'event: done' in r.text
    assert 'grounded": true' in r.text
    assert 'chat_key' not in client.get('/api/v1/settings').json()


def test_backup_restore_and_tamper_rejection(client):
    m=create(client,title='备份保留')
    source=process_upload(client,m['id'],'backup.txt','备份原件'.encode())
    archive=client.get('/api/v1/backups/export')
    assert archive.status_code==200,archive.text[:100]
    create(client,title='备份后的内容')
    original_path=main.library.path
    r=client.post('/api/v1/backups/restore',files={'file':('backup.zip',archive.content)})
    assert r.status_code==200,r.text
    assert original_path.exists() and main.library.path!=original_path
    assert len(client.get('/api/v1/entries?kind=manual').json()['items'])==1
    assert client.get('/api/v1/blobs/'+source['attrs']['hash']).content=='备份原件'.encode()
    bad=io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(archive.content)) as old,zipfile.ZipFile(bad,'w') as new:
        for name in old.namelist():
            raw=old.read(name)
            if name.startswith('blobs/'):
                raw=b'x'*len(raw)
            new.writestr(name,raw)
    path=main.library.path
    r=client.post('/api/v1/backups/restore',files={'file':('bad.zip',bad.getvalue())})
    assert r.status_code==400 and main.library.path==path


def test_security_and_missing_ai(client):
    assert client.get('/api/v1/settings',headers={'Origin':'https://evil.invalid'}).status_code==403
    assert client.get('/api/v1/entries',headers={'Host':'evil.invalid'}).status_code==403
    for url in ['http://127.0.0.1/private','file:///etc/passwd','http://user@example.com/']:
        with pytest.raises(ValueError):safe_url(url)
    assert client.post('/api/v1/chat',json={'question':'你好'}).status_code==400
    if os.name=='nt':
        assert protect(protect(b'test secret'),True)==b'test secret'


def test_processing_rechecks_deleted_before_commit(client,monkeypatch):
    m=create(client)
    r=client.post('/api/v1/imports/file',data={'manual_id':m['id']},files={'file':('a.txt',b'original')}).json()
    sid=r['source']['id']
    client.post('/api/v1/entries/'+sid+'/trash',json={'revision':1})
    with pytest.raises(ValueError):main.processor.process(r['job_id'],sid)
    assert search(client,'original')==[]


def test_organize_review_apply_and_conflict(client, monkeypatch):
    m=create(client,title='待整理',attrs={'ai_enabled':True})
    create(client,kind='note',manual_id=m['id'],content='UART 工作电压 3.3V')
    from backend import organization
    monkeypatch.setattr(organization, 'chat_json', lambda *args, **kwargs: {
        'chapters': ['接口', '使用', '接口'], 'summary': '根据个人记录整理',
        'cards': [{'chapter': '接口', 'title': 'UART', 'content': '根据个人记录：工作电压 3.3V', 'source_ref': 1}],
    })
    proposal=client.post('/api/v1/manuals/'+m['id']+'/organize')
    assert proposal.status_code==200,proposal.text
    assert client.get('/api/v1/entries?kind=chapter').json()['items']==[]
    r=client.post('/api/v1/manuals/'+m['id']+'/organize/apply',json=proposal.json())
    assert r.status_code==200,r.text
    assert len(client.get('/api/v1/entries?kind=chapter').json()['items'])==2
    assert client.post('/api/v1/manuals/'+m['id']+'/organize/apply',json=proposal.json()).status_code==409


def test_source_correction_keeps_original_and_reindexes(client):
    m=create(client)
    source=process_upload(client,m['id'],'text.txt',b'original 33V')
    path='/api/v1/sources/'+source['id']+'/text'
    payload={'revision':source['revision'],'block':0,'text':'corrected 3.3V'}
    assert client.patch(path,json=payload).status_code==200
    assert client.patch(path,json=payload).status_code==409
    assert search(client,'corrected') and not search(client,'original')
    block=client.get('/api/v1/entries/'+source['id']).json()['attrs']['blocks'][0]
    assert block['original_text']=='original 33V'
    assert client.get('/api/v1/blobs/'+source['attrs']['hash']).content==b'original 33V'


def test_category_rename_trash_restore(client):
    m=create(client,category='电子')
    cat=client.get('/api/v1/entries?kind=category').json()['items'][0]
    r=client.patch('/api/v1/entries/'+cat['id'],json={'revision':1,'title':'硬件'})
    assert r.status_code==200
    assert client.get('/api/v1/entries/'+m['id']).json()['category']=='硬件'
    assert client.post('/api/v1/entries/'+cat['id']+'/trash',json={'revision':2}).status_code==200
    assert client.get('/api/v1/entries/'+m['id']).json()['category']==''
    assert client.post('/api/v1/entries/'+cat['id']+'/restore').status_code==200
    assert client.get('/api/v1/entries/'+m['id']).json()['category']=='硬件'


def test_real_text_and_scanned_pdf(client):
    from PIL import Image,ImageDraw,ImageFont
    m=create(client)
    stream=b'BT /F1 16 Tf 50 700 Td (UART voltage 3.3V. GPIO18 installation instructions.) Tj ET'
    objects=[b'<< /Type /Catalog /Pages 2 0 R >>',b'<< /Type /Pages /Kids [3 0 R] /Count 1 >>',b'<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>',b'<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>',b'<< /Length '+str(len(stream)).encode()+b' >>\nstream\n'+stream+b'\nendstream']
    raw=bytearray(b'%PDF-1.4\n');offsets=[]
    for i,obj in enumerate(objects,1):
        offsets.append(len(raw));raw.extend(str(i).encode()+b' 0 obj\n'+obj+b'\nendobj\n')
    xref=len(raw);raw.extend(b'xref\n0 6\n0000000000 65535 f \n')
    for offset in offsets:raw.extend(f'{offset:010d} 00000 n \n'.encode())
    raw.extend(b'trailer\n<< /Size 6 /Root 1 0 R >>\nstartxref\n'+str(xref).encode()+b'\n%%EOF')
    source=process_upload(client,m['id'],'text.pdf',bytes(raw),'application/pdf')
    assert '3.3V' in source['attrs']['blocks'][0]['text']
    image=Image.new('RGB',(1200,450),'white');draw=ImageDraw.Draw(image)
    font=ImageFont.truetype('C:/Windows/Fonts/arial.ttf',48)
    draw.text((40,80),'UART voltage 3.3V',fill='black',font=font)
    scanned=io.BytesIO();image.save(scanned,format='PDF')
    source=process_upload(client,m['id'],'scanned.pdf',scanned.getvalue(),'application/pdf')
    assert source['attrs']['blocks'][0]['locator']['ocr']
    assert 'UART' in source['attrs']['blocks'][0]['text']


def test_parser_subprocess(client):
    m=create(client)
    r=client.post('/api/v1/imports/file',data={'manual_id':m['id']},files={'file':('worker.txt',b'isolated process result')}).json()
    main.processor.thread=object()
    try:main.processor.process(r['job_id'],r['source']['id'])
    finally:main.processor.thread=None
    assert search(client,'isolated')


def test_real_http_ai_adapter_and_index(client):
    from http.server import BaseHTTPRequestHandler,ThreadingHTTPServer
    import threading
    received=[]
    class Provider(BaseHTTPRequestHandler):
        def log_message(self,*args):pass
        def do_POST(self):
            payload=json.loads(self.rfile.read(int(self.headers['Content-Length'])))
            received.append((self.path,payload))
            if self.path.endswith('/embeddings'):
                body=json.dumps({'data':[{'index':i,'embedding':[1,0,0]} for i,_ in enumerate(payload['input'])]}).encode()
                content_type='application/json'
            else:
                body=('data: '+json.dumps({'choices':[{'delta':{'content':'UART 电压为 3.3V [1]。'}}]})+'\n\ndata: [DONE]\n\n').encode()
                content_type='text/event-stream'
            self.send_response(200);self.send_header('Content-Type',content_type);self.send_header('Content-Length',str(len(body)));self.end_headers();self.wfile.write(body)
    server=ThreadingHTTPServer(('127.0.0.1',0),Provider)
    thread=threading.Thread(target=server.serve_forever,daemon=True);thread.start()
    try:
        base=f'http://127.0.0.1:{server.server_port}/v1'
        r=client.put('/api/v1/settings',json={'chat_base':base,'chat_model':'fixture','chat_key':'fixture-only','embedding_base':base,'embedding_model':'fixture','embedding_key':'fixture-only'})
        assert r.status_code==200,r.text
        m=create(client,attrs={'ai_enabled':True})
        create(client,kind='card',manual_id=m['id'],title='UART 电压',content='UART 工作电压 3.3V')
        with main.library.Session() as db:jid=db.scalar(select(Job.id).where(Job.source_id==m['id'],Job.kind=='index'))
        main.processor.index(jid,m['id'])
        assert client.post('/api/v1/search',json={'q':'UART 电压','mode':'semantic'}).json()['items']
        r=client.post('/api/v1/chat',json={'question':'UART 电压','manual_id':m['id']})
        assert 'grounded": true' in r.text and 'event: delta' in r.text
        assert any(path.endswith('/chat/completions') for path,_ in received)
        assert 'fixture-only' not in client.get('/api/v1/settings').text
    finally:server.shutdown();server.server_close();thread.join(timeout=3)
